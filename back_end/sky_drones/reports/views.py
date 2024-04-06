import io
import uuid
from datetime import datetime

import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from rest_framework import permissions, status
from rest_framework.response import Response
from rest_framework.views import APIView

from app_users.models import AppUser
from defects.models import Defect
from facilities.models import Facility
from file_storage_items.models import FileStorageItem, Template
from file_storage_items.utils import FileStorageItemType
from inspections.models import Inspection
from reports.models import Report
from sky_drones.mixins import AmazonS3Mixin
from sky_drones.utils import RoleEmployeeBasedPermission

matplotlib.use('Agg')


class ReportUrlAPIRetrieve(APIView, AmazonS3Mixin):
    permission_classes = (permissions.IsAuthenticated,)

    def get(self, request, inspection_id):
        try:
            report = Report.objects.get(inspection_id=inspection_id)
            file_storage_item = FileStorageItem.objects.get(id=report.file_storage_item_id)
            presigned_report_url = self.get_generated_presigned_url(file_storage_item.file_name)
        except Report.DoesNotExist:
            presigned_report_url = None
        except FileStorageItem.DoesNotExist:
            presigned_report_url = None

        return Response({'data': presigned_report_url}, status=status.HTTP_200_OK)


class ReportAPICreate(APIView, AmazonS3Mixin):
    permission_classes = (permissions.IsAuthenticated, RoleEmployeeBasedPermission,)

    def post(self, request):
        author = self.request.user
        created_date = datetime.now().date()
        inspection_id = request.data.get('inspectionId')
        inspection = Inspection.objects.get(id=inspection_id)
        facility = Facility.objects.get(id=inspection.facility.id)
        pilot = AppUser.objects.get(id=inspection.pilot.id)
        inspector = AppUser.objects.get(id=inspection.inspector.id)
        defects_queryset = Defect.objects.filter(inspection_id=inspection_id)
        file_storage_item_template_id = 11
        template_id = 1

        file_storage_item = FileStorageItem.objects.get(id=file_storage_item_template_id)
        temp_template_path = self.download_file_from_s3(file_storage_item.file_name)
        if not temp_template_path:
            return Response({'detail': 'Error downloading template from S3'},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        variables = {
            "{created_date}": created_date.strftime("%Y-%m-%d"),
            "{author_username}": author.username,
            "{facility_name}": facility.name,
            "{inspection_name}": inspection.name,
            "{pilot_username}": pilot.username,
            "{inspector_username}": inspector.username,
            "{reason}": inspection.reason,
        }

        document = Document(temp_template_path)

        for variable_key, variable_value in variables.items():
            for table in document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        buffer = create_severity_pie_diagram_buffer(defects_queryset)
        document.add_picture(buffer, width=Inches(6.5))

        style = document.styles['Normal']
        font = style.font
        font.size = Pt(16)

        red = RGBColor(194, 16, 3)
        orange = RGBColor(237, 118, 36)
        yellow = RGBColor(255, 201, 14)
        blue = RGBColor(0, 153, 219)

        severity_colors = {
            "CRITICAL": red,
            "MAJOR": orange,
            "MINOR": yellow,
            "TRIVIAL": blue
        }

        for defect in defects_queryset:
            file_storage_item = FileStorageItem.objects.get(id=defect.file_storage_item_id)
            temp_image_path = self.download_file_from_s3(file_storage_item.file_name)
            if not temp_image_path:
                return Response({'detail': 'Error downloading image from S3'},
                                status=status.HTTP_500_INTERNAL_SERVER_ERROR)

            document.add_page_break()
            document.add_heading(defect.name, level=1)
            paragraph = document.add_paragraph()
            paragraph.add_run("Severity: ")
            paragraph.add_run(f"{defect.severity}").font.color.rgb = severity_colors.get(defect.severity)
            document.add_paragraph(f"Description: {defect.description}")
            document.add_paragraph(f"")
            document.add_picture(temp_image_path, width=Inches(6.5))

        document_buffer = io.BytesIO()
        document.save(document_buffer)
        document_buffer.seek(0)

        unique_filename = str(uuid.uuid4()) + ".docx"
        url = self.upload_to_s3(document_buffer, unique_filename)
        if not url:
            return Response({'detail': 'Failed to upload report to S3'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        file_storage_item = FileStorageItem.objects.create(
            file_name=unique_filename,
            type=FileStorageItemType.REPORT.value,
            path=url,
            facility=facility
        )

        template = Template.objects.get(id=template_id)
        Report.objects.create(
            created_date=datetime.now(),
            author=author,
            template=template,
            inspection=inspection,
            file_storage_item=file_storage_item
        )

        presigned_url = self.get_generated_presigned_url(file_storage_item.file_name)

        return Response({'presigned_url': presigned_url}, status=status.HTTP_201_CREATED)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


def create_severity_pie_diagram_buffer(defects_queryset):
    severity_counts = {
        "CRITICAL": 0,
        "MAJOR": 0,
        "MINOR": 0,
        "TRIVIAL": 0
    }
    for defect in defects_queryset:
        severity_counts[defect.severity] += 1

    sizes = [severity_counts["CRITICAL"],
             severity_counts["MAJOR"],
             severity_counts["MINOR"],
             severity_counts["TRIVIAL"]]

    labels = ['CRITICAL', 'MAJOR', 'MINOR', 'TRIVIAL']
    colors = ['red', 'orange', 'yellow', 'lightskyblue']

    plt.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
    plt.title('Defects severity')

    buffer = io.BytesIO()
    plt.savefig(buffer, format='png')
    buffer.seek(0)
    return buffer
